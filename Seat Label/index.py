import docx
from docx.shared import Pt,Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
import pandas as pd
reg = pd.read_csv('Copy of 18th AAM - Registration Panel - new main.csv',header=None)


mydoc = docx.Document()
section = mydoc.sections[0]
section.page_height = Cm(21)
section.page_width = Cm(29.7)
section.orientation = WD_ORIENT.LANDSCAPE
 
rk = ['RK','R K', 'Radhakrishnan Hall of Residence','RK Hall','R K Hall']

newpage = 0

for i in range(1,len(reg)):
    doc = mydoc.add_paragraph()
    doc.add_run(reg[3][i]).font.size = Pt(60)
    doc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc = mydoc.add_paragraph()
    hall = ""
    if str(reg[31][i]).upper()[0] == 'R':
        if str(reg[31][i]) in rk:
            hall = 'RK'
        else:
            hall = 'RP' 
    elif str(reg[31][i]).upper()[0] == 'V':
        hall = 'VS'
    elif str(reg[31][i]).upper()[0] == 'J':
        hall = 'JCB'
    elif str(reg[31][i]).upper()[0] == 'A':
        hall = 'Azad'
    elif str(reg[31][i]).upper()[0] == 'P':
        hall = 'Patel'
    elif str(reg[31][i]).upper()[0] == 'N':
        hall = 'Nehru'
    elif str(reg[31][i]).upper()[0] == 'L':
        hall = 'LLR'
    elif str(reg[31][i]).upper()[0] == 'S':
        hall = 'SN/IG'
    elif str(reg[31][i]).upper()[0] == 'Z':
        hall = 'ZRH'
    
    doc.add_run(str(reg[32][i])+" | "+hall+" | <DEP>").font.size = Pt(54)
    doc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mydoc.add_paragraph().add_run('\n').font.size = Pt(30)
    if newpage%2 == 1:
        mydoc.add_page_break()

    newpage+=1


print(mydoc.sections[0].orientation)

mydoc.save("out.docx")